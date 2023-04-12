import docx
from scout import mysql_scout2 as scrape
from docx.shared import Cm

def generatereport(scoutteam, season, playerstats, teampaint):

    document = docx.Document('/home/emiller/d2east-django/scout/template' + str(len(playerstats)) + '.docx')

    headerpar = document.paragraphs[0]
    headertext = headerpar._element.xpath('.//w:t')
    p1par = document.paragraphs[4]
    p1text = p1par._element.xpath('.//w:t')
    p2par = document.paragraphs[15]
    p2text = p2par._element.xpath('.//w:t')
    if len(playerstats) == 7 or len(playerstats) == 6:
        teampar = document.paragraphs[29]
    else:
        teampar = document.paragraphs[27]
    teamtext = teampar._element.xpath('.//w:t')

    teamstats = scrape.getteamstats(scoutteam, season)
    oppstats = scrape.getoppstats(scoutteam, season)

    gp = int(teamstats[2])
    fgm = int(teamstats[22])
    fga = int(teamstats[23])
    threem = int(teamstats[20])
    threea = int(teamstats[21])
    if teampaint[0] == 0:
        paintpct = 0
    else:
        paintpct = round(teampaint[0] / teampaint[1] * 100)
    midm = fgm - threem - teampaint[0]
    mida = fga - threea - teampaint[1]
    midmg = round(midm/gp,1)
    midag = round(mida/gp,1)
    paintmg = round(teampaint[0]/gp,1)
    paintag = round(teampaint[1]/gp,1)
    if midm == 0:
        midpct = 0
    else:
        midpct = round((midm/mida)*100)
    paintdist = round(teampaint[1]/fga * 100)
    middist = round(mida/fga * 100)
    threedist = round(threea/fga * 100)

    player1stat = scrape.scoutplayer(scoutteam, season, playerstats[0][0], playerstats[0][1], playerstats[0][2])
    player2stat = scrape.scoutplayer(scoutteam, season, playerstats[1][0], playerstats[1][1], playerstats[1][2])
    player3stat = scrape.scoutplayer(scoutteam, season, playerstats[2][0], playerstats[2][1], playerstats[2][2])
    player4stat = scrape.scoutplayer(scoutteam, season, playerstats[3][0], playerstats[3][1], playerstats[3][2])
    player5stat = scrape.scoutplayer(scoutteam, season, playerstats[4][0], playerstats[4][1], playerstats[4][2])
    player6stat = scrape.scoutplayer(scoutteam, season, playerstats[5][0], playerstats[5][1], playerstats[5][2])
    if len(playerstats) >= 7:
        player7stat = scrape.scoutplayer(scoutteam, season, playerstats[6][0], playerstats[6][1], playerstats[6][2])
    if len(playerstats) >= 8:
        player8stat = scrape.scoutplayer(scoutteam, season, playerstats[7][0], playerstats[7][1], playerstats[7][2])
    if len(playerstats) >= 9:
        player9stat = scrape.scoutplayer(scoutteam, season, playerstats[8][0], playerstats[8][1], playerstats[8][2])
    if len(playerstats) >= 10:
        player10stat = scrape.scoutplayer(scoutteam, season, playerstats[9][0], playerstats[9][1], playerstats[9][2])
    if len(playerstats) >= 11:
        player11stat = scrape.scoutplayer(scoutteam, season, playerstats[10][0], playerstats[10][1], playerstats[10][2])

    headerinfo = scrape.headerinfo(scoutteam, season)

    headertext[0].text = str(scrape.getteamnamefromabbr(scoutteam))
    headertext[5].text = str(headerinfo[0])
    headertext[7].text = str(headerinfo[1])
    headertext[9].text = str(headerinfo[2])
    headertext[11].text = str(headerinfo[3])
    headertext[13].text = str(scrape.getteamconffromabbr(scoutteam))

    img = document.tables[0].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/teamicons/' + scoutteam + '.png', height=Cm(2.75))

    # first player

    img = document.tables[1].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/playerphotos/' + player1stat[0] + '.jpg', height=Cm(3.39))

    p1text[1].text = str(player1stat[1])
    p1text[2].text = str(player1stat[2])
    p1text[3].text = str(player1stat[3])
    p1text[4].text = str(player1stat[4])
    p1text[5].text = str(player1stat[5])
    p1text[6].text = str(player1stat[6])
    p1text[8].text = str(player1stat[7])
    p1text[10].text = str(player1stat[8])
    p1text[15].text = str(player1stat[9])
    p1text[18].text = str(player1stat[10])
    p1text[20].text = str(player1stat[11])
    p1text[22].text = str(player1stat[12])
    p1text[26].text = str(player1stat[13])
    p1text[28].text = str(player1stat[14])
    p1text[30].text = str(player1stat[15])
    p1text[34].text = str(player1stat[16])
    p1text[36].text = str(player1stat[17])
    p1text[38].text = str(player1stat[18])
    p1text[44].text = str(player1stat[19])
    p1text[46].text = str(player1stat[20])
    p1text[49].text = str(player1stat[21])
    p1text[51].text = str(player1stat[22])
    p1text[54].text = str(player1stat[23])
    p1text[56].text = str(player1stat[24])
    #
    p1text[59].text = str(player1stat[25])
    p1text[61].text = str(player1stat[26])
    p1text[63].text = str(player1stat[27])
    p1text[67].text = str(player1stat[28])
    p1text[69].text = str(player1stat[29])
    p1text[71].text = str(player1stat[30])
    p1text[75].text = str(player1stat[31])
    p1text[77].text = str(player1stat[32])
    p1text[79].text = str(player1stat[33])

    p1text[83].text = ''

    # second player

    img = document.tables[2].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/playerphotos/' + player2stat[0] + '.jpg', height=Cm(3.39))

    p1text[85].text = str(player2stat[1])
    p1text[86].text = str(player2stat[2])
    p1text[87].text = str(player2stat[3])
    p1text[88].text = str(player2stat[4])
    p1text[89].text = str(player2stat[5])
    p1text[90].text = str(player2stat[6])
    p1text[92].text = str(player2stat[7])
    p1text[94].text = str(player2stat[8])
    p1text[99].text = str(player2stat[9])
    p1text[102].text = str(player2stat[10])
    p1text[104].text = str(player2stat[11])
    p1text[106].text = str(player2stat[12])
    p1text[110].text = str(player2stat[13])
    p1text[112].text = str(player2stat[14])
    p1text[114].text = str(player2stat[15])
    p1text[118].text = str(player2stat[16])
    p1text[120].text = str(player2stat[17])
    p1text[122].text = str(player2stat[18])
    p1text[128].text = str(player2stat[19])
    p1text[130].text = str(player2stat[20])
    p1text[133].text = str(player2stat[21])
    p1text[135].text = str(player2stat[22])
    p1text[138].text = str(player2stat[23])
    p1text[140].text = str(player2stat[24])
    #
    p1text[143].text = str(player2stat[25])
    p1text[145].text = str(player2stat[26])
    p1text[147].text = str(player2stat[27])
    p1text[151].text = str(player2stat[28])
    p1text[153].text = str(player2stat[29])
    p1text[155].text = str(player2stat[30])
    p1text[159].text = str(player2stat[31])
    p1text[161].text = str(player2stat[32])
    p1text[163].text = str(player2stat[33])

    p1text[167].text = ''

    # third player

    img = document.tables[3].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/playerphotos/' + player3stat[0] + '.jpg', height=Cm(3.39))

    p1text[169].text = str(player3stat[1])
    p1text[170].text = str(player3stat[2])
    p1text[171].text = str(player3stat[3])
    p1text[172].text = str(player3stat[4])
    p1text[173].text = str(player3stat[5])
    p1text[174].text = str(player3stat[6])
    p1text[176].text = str(player3stat[7])
    p1text[178].text = str(player3stat[8])
    p1text[183].text = str(player3stat[9])
    p1text[186].text = str(player3stat[10])
    p1text[188].text = str(player3stat[11])
    p1text[190].text = str(player3stat[12])
    p1text[194].text = str(player3stat[13])
    p1text[196].text = str(player3stat[14])
    p1text[198].text = str(player3stat[15])
    p1text[202].text = str(player3stat[16])
    p1text[204].text = str(player3stat[17])
    p1text[206].text = str(player3stat[18])
    p1text[212].text = str(player3stat[19])
    p1text[214].text = str(player3stat[20])
    p1text[217].text = str(player3stat[21])
    p1text[219].text = str(player3stat[22])
    p1text[222].text = str(player3stat[23])
    p1text[224].text = str(player3stat[24])
    #
    p1text[227].text = str(player3stat[25])
    p1text[229].text = str(player3stat[26])
    p1text[231].text = str(player3stat[27])
    p1text[235].text = str(player3stat[28])
    p1text[237].text = str(player3stat[29])
    p1text[239].text = str(player3stat[30])
    p1text[243].text = str(player3stat[31])
    p1text[245].text = str(player3stat[32])
    p1text[247].text = str(player3stat[33])

    p1text[251].text = '' # +4

    # fourth player

    img = document.tables[4].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/playerphotos/' + player4stat[0] + '.jpg', height=Cm(3.39))

    p1text[253].text = str(player4stat[1])
    p1text[254].text = str(player4stat[2])
    p1text[255].text = str(player4stat[3])
    p1text[256].text = str(player4stat[4])
    p1text[257].text = str(player4stat[5])
    p1text[258].text = str(player4stat[6])
    p1text[260].text = str(player4stat[7])
    p1text[262].text = str(player4stat[8])
    p1text[267].text = str(player4stat[9])
    p1text[270].text = str(player4stat[10])
    p1text[272].text = str(player4stat[11])
    p1text[274].text = str(player4stat[12])
    p1text[278].text = str(player4stat[13])
    p1text[280].text = str(player4stat[14])
    p1text[282].text = str(player4stat[15])
    p1text[286].text = str(player4stat[16])
    p1text[288].text = str(player4stat[17])
    p1text[290].text = str(player4stat[18])
    p1text[296].text = str(player4stat[19])
    p1text[298].text = str(player4stat[20])
    p1text[301].text = str(player4stat[21])
    p1text[303].text = str(player4stat[22])
    p1text[306].text = str(player4stat[23])
    p1text[308].text = str(player4stat[24])
    #
    p1text[311].text = str(player4stat[25])
    p1text[313].text = str(player4stat[26])
    p1text[315].text = str(player4stat[27])
    p1text[319].text = str(player4stat[28])
    p1text[321].text = str(player4stat[29])
    p1text[323].text = str(player4stat[30])
    p1text[327].text = str(player4stat[31])
    p1text[329].text = str(player4stat[32])
    p1text[331].text = str(player4stat[33])

    p1text[335].text = ''

    # fifth player

    img = document.tables[5].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/playerphotos/' + player5stat[0] + '.jpg', height=Cm(3.39))

    p1text[337].text = str(player5stat[1])
    p1text[338].text = str(player5stat[2]) # +1
    p1text[339].text = str(player5stat[3]) # +1
    p1text[340].text = str(player5stat[4]) # +1
    p1text[341].text = str(player5stat[5]) # +1
    p1text[342].text = str(player5stat[6]) # +1
    p1text[344].text = str(player5stat[7]) # +2
    p1text[346].text = str(player5stat[8]) # +2
    p1text[351].text = str(player5stat[9]) # +5
    p1text[354].text = str(player5stat[10]) # +3
    p1text[356].text = str(player5stat[11]) # +2
    p1text[358].text = str(player5stat[12]) # +2
    p1text[362].text = str(player5stat[13]) # +4
    p1text[364].text = str(player5stat[14]) # +2
    p1text[366].text = str(player5stat[15]) # +2
    p1text[370].text = str(player5stat[16]) # +4
    p1text[372].text = str(player5stat[17]) # +2
    p1text[374].text = str(player5stat[18]) # +2
    p1text[380].text = str(player5stat[19]) # +6
    p1text[382].text = str(player5stat[20]) # +2
    p1text[385].text = str(player5stat[21]) # +3
    p1text[387].text = str(player5stat[22]) # +2
    p1text[390].text = str(player5stat[23]) # +3
    p1text[392].text = str(player5stat[24]) # +2
    #
    p1text[395].text = str(player5stat[25]) # +3
    p1text[397].text = str(player5stat[26]) # +2
    p1text[399].text = str(player5stat[27]) # +2
    p1text[403].text = str(player5stat[28]) # +4
    p1text[405].text = str(player5stat[29]) # +2
    p1text[407].text = str(player5stat[30]) # +2
    p1text[411].text = str(player5stat[31]) # +4
    p1text[413].text = str(player5stat[32]) # +2
    p1text[415].text = str(player5stat[33]) # +2

    p1text[419].text = '' # +4

    # first bench player

    img = document.tables[6].rows[0].cells[0].add_paragraph()
    run = img.add_run()
    run.add_picture('/home/emiller/d2east-django/playerphotos/' + player6stat[0] + '.jpg', height=Cm(3.39))

    p2text[1].text = str(player6stat[1])
    p2text[2].text = str(player6stat[2])
    p2text[3].text = str(player6stat[3])
    p2text[4].text = str(player6stat[4])
    p2text[5].text = str(player6stat[5])
    p2text[6].text = str(player6stat[6])
    p2text[8].text = str(player6stat[7])
    p2text[10].text = str(player6stat[8])
    p2text[15].text = str(player6stat[9])
    p2text[18].text = str(player6stat[10])
    p2text[20].text = str(player6stat[11])
    p2text[22].text = str(player6stat[12])
    p2text[26].text = str(player6stat[13])
    p2text[28].text = str(player6stat[14])
    p2text[30].text = str(player6stat[15])
    p2text[34].text = str(player6stat[16])
    p2text[36].text = str(player6stat[17])
    p2text[38].text = str(player6stat[18])
    p2text[44].text = str(player6stat[19])
    p2text[46].text = str(player6stat[20])
    p2text[49].text = str(player6stat[21])
    p2text[51].text = str(player6stat[22])
    p2text[54].text = str(player6stat[23])
    p2text[56].text = str(player6stat[24])
    #
    p2text[59].text = str(player6stat[25])
    p2text[61].text = str(player6stat[26])
    p2text[63].text = str(player6stat[27])
    p2text[67].text = str(player6stat[28])
    p2text[69].text = str(player6stat[29])
    p2text[71].text = str(player6stat[30])
    p2text[75].text = str(player6stat[31])
    p2text[77].text = str(player6stat[32])
    p2text[79].text = str(player6stat[33])

    p2text[83].text = ''

    if len(playerstats) >= 7:
        # second bench player

        img = document.tables[7].rows[0].cells[0].add_paragraph()
        run = img.add_run()
        run.add_picture('/home/emiller/d2east-django/playerphotos/' + player7stat[0] + '.jpg', height=Cm(3.39))

        p2text[85].text = str(player7stat[1])
        p2text[86].text = str(player7stat[2])
        p2text[87].text = str(player7stat[3])
        p2text[88].text = str(player7stat[4])
        p2text[89].text = str(player7stat[5])
        p2text[90].text = str(player7stat[6])
        p2text[92].text = str(player7stat[7])
        p2text[94].text = str(player7stat[8])
        p2text[99].text = str(player7stat[9])
        p2text[102].text = str(player7stat[10])
        p2text[104].text = str(player7stat[11])
        p2text[106].text = str(player7stat[12])
        p2text[110].text = str(player7stat[13])
        p2text[112].text = str(player7stat[14])
        p2text[114].text = str(player7stat[15])
        p2text[118].text = str(player7stat[16])
        p2text[120].text = str(player7stat[17])
        p2text[122].text = str(player7stat[18])
        p2text[128].text = str(player7stat[19])
        p2text[130].text = str(player7stat[20])
        p2text[133].text = str(player7stat[21])
        p2text[135].text = str(player7stat[22])
        p2text[138].text = str(player7stat[23])
        p2text[140].text = str(player7stat[24])
        #
        p2text[143].text = str(player7stat[25])
        p2text[145].text = str(player7stat[26])
        p2text[147].text = str(player7stat[27])
        p2text[151].text = str(player7stat[28])
        p2text[153].text = str(player7stat[29])
        p2text[155].text = str(player7stat[30])
        p2text[159].text = str(player7stat[31])
        p2text[161].text = str(player7stat[32])
        p2text[163].text = str(player7stat[33])

        p2text[167].text = ''

    if len(playerstats) >= 8:

        # third bench player

        img = document.tables[8].rows[0].cells[0].add_paragraph()
        run = img.add_run()
        run.add_picture('/home/emiller/d2east-django/playerphotos/' + player8stat[0] + '.jpg', height=Cm(3.39))

        p2text[169].text = str(player8stat[1])
        p2text[170].text = str(player8stat[2])
        p2text[171].text = str(player8stat[3])
        p2text[172].text = str(player8stat[4])
        p2text[173].text = str(player8stat[5])
        p2text[174].text = str(player8stat[6])
        p2text[176].text = str(player8stat[7])
        p2text[178].text = str(player8stat[8])
        p2text[183].text = str(player8stat[9])
        p2text[186].text = str(player8stat[10])
        p2text[188].text = str(player8stat[11])
        p2text[190].text = str(player8stat[12])
        p2text[194].text = str(player8stat[13])
        p2text[196].text = str(player8stat[14])
        p2text[198].text = str(player8stat[15])
        p2text[202].text = str(player8stat[16])
        p2text[204].text = str(player8stat[17])
        p2text[206].text = str(player8stat[18])
        p2text[212].text = str(player8stat[19])
        p2text[214].text = str(player8stat[20])
        p2text[217].text = str(player8stat[21])
        p2text[219].text = str(player8stat[22])
        p2text[222].text = str(player8stat[23])
        p2text[224].text = str(player8stat[24])
        #
        p2text[227].text = str(player8stat[25])
        p2text[229].text = str(player8stat[26])
        p2text[231].text = str(player8stat[27])
        p2text[235].text = str(player8stat[28])
        p2text[237].text = str(player8stat[29])
        p2text[239].text = str(player8stat[30])
        p2text[243].text = str(player8stat[31])
        p2text[245].text = str(player8stat[32])
        p2text[247].text = str(player8stat[33])

        p2text[251].text = '' # +4

    if len(playerstats) >= 9:

        # fourth bench player

        img = document.tables[9].rows[0].cells[0].add_paragraph()
        run = img.add_run()
        run.add_picture('/home/emiller/d2east-django/playerphotos/' + player9stat[0] + '.jpg', height=Cm(3.39))

        p2text[253].text = str(player9stat[1])
        p2text[254].text = str(player9stat[2])
        p2text[255].text = str(player9stat[3])
        p2text[256].text = str(player9stat[4])
        p2text[257].text = str(player9stat[5])
        p2text[258].text = str(player9stat[6])
        p2text[260].text = str(player9stat[7])
        p2text[262].text = str(player9stat[8])
        p2text[267].text = str(player9stat[9])
        p2text[270].text = str(player9stat[10])
        p2text[272].text = str(player9stat[11])
        p2text[274].text = str(player9stat[12])
        p2text[278].text = str(player9stat[13])
        p2text[280].text = str(player9stat[14])
        p2text[282].text = str(player9stat[15])
        p2text[286].text = str(player9stat[16])
        p2text[288].text = str(player9stat[17])
        p2text[290].text = str(player9stat[18])
        p2text[296].text = str(player9stat[19])
        p2text[298].text = str(player9stat[20])
        p2text[301].text = str(player9stat[21])
        p2text[303].text = str(player9stat[22])
        p2text[306].text = str(player9stat[23])
        p2text[308].text = str(player9stat[24])
        #
        p2text[311].text = str(player9stat[25])
        p2text[313].text = str(player9stat[26])
        p2text[315].text = str(player9stat[27])
        p2text[319].text = str(player9stat[28])
        p2text[321].text = str(player9stat[29])
        p2text[323].text = str(player9stat[30])
        p2text[327].text = str(player9stat[31])
        p2text[329].text = str(player9stat[32])
        p2text[331].text = str(player9stat[33])

        p2text[335].text = ''

    if len(playerstats) >= 10:

        # fifth bench player

        img = document.tables[10].rows[0].cells[0].add_paragraph()
        run = img.add_run()
        run.add_picture('/home/emiller/d2east-django/playerphotos/' + player10stat[0] + '.jpg', height=Cm(3.39))

        p2text[337].text = str(player10stat[1])
        p2text[338].text = str(player10stat[2])
        p2text[339].text = str(player10stat[3])
        p2text[340].text = str(player10stat[4])
        p2text[341].text = str(player10stat[5])
        p2text[342].text = str(player10stat[6])
        p2text[344].text = str(player10stat[7])
        p2text[346].text = str(player10stat[8])
        p2text[351].text = str(player10stat[9])
        p2text[354].text = str(player10stat[10])
        p2text[356].text = str(player10stat[11])
        p2text[358].text = str(player10stat[12])
        p2text[362].text = str(player10stat[13])
        p2text[364].text = str(player10stat[14])
        p2text[366].text = str(player10stat[15])
        p2text[370].text = str(player10stat[16])
        p2text[372].text = str(player10stat[17])
        p2text[374].text = str(player10stat[18])
        p2text[380].text = str(player10stat[19])
        p2text[382].text = str(player10stat[20])
        p2text[385].text = str(player10stat[21])
        p2text[387].text = str(player10stat[22])
        p2text[390].text = str(player10stat[23])
        p2text[392].text = str(player10stat[24])
        #
        p2text[395].text = str(player10stat[25])
        p2text[397].text = str(player10stat[26])
        p2text[399].text = str(player10stat[27])
        p2text[403].text = str(player10stat[28])
        p2text[405].text = str(player10stat[29])
        p2text[407].text = str(player10stat[30])
        p2text[411].text = str(player10stat[31])
        p2text[413].text = str(player10stat[32])
        p2text[415].text = str(player10stat[33])

        p2text[419].text = '' # +4

    if len(playerstats) >= 11:

        # sixth bench player

        img = document.tables[11].rows[0].cells[0].add_paragraph()
        run = img.add_run()
        run.add_picture('/home/emiller/d2east-django/playerphotos/' + player11stat[0] + '.jpg', height=Cm(3.39))

        p2text[421].text = str(player11stat[1])
        p2text[422].text = str(player11stat[2])
        p2text[423].text = str(player11stat[3])
        p2text[424].text = str(player11stat[4])
        p2text[425].text = str(player11stat[5])
        p2text[426].text = str(player11stat[6])
        p2text[428].text = str(player11stat[7])
        p2text[430].text = str(player11stat[8])
        p2text[435].text = str(player11stat[9])
        p2text[438].text = str(player11stat[10])
        p2text[440].text = str(player11stat[11])
        p2text[442].text = str(player11stat[12])
        p2text[446].text = str(player11stat[13])
        p2text[448].text = str(player11stat[14])
        p2text[450].text = str(player11stat[15])
        p2text[454].text = str(player11stat[16])
        p2text[456].text = str(player11stat[17])
        p2text[458].text = str(player11stat[18])
        p2text[464].text = str(player11stat[19])
        p2text[466].text = str(player11stat[20])
        p2text[469].text = str(player11stat[21])
        p2text[471].text = str(player11stat[22])
        p2text[474].text = str(player11stat[23])
        p2text[476].text = str(player11stat[24])
        #
        p2text[479].text = str(player11stat[25])
        p2text[481].text = str(player11stat[26])
        p2text[483].text = str(player11stat[27])
        p2text[487].text = str(player11stat[28])
        p2text[489].text = str(player11stat[29])
        p2text[491].text = str(player11stat[30])
        p2text[495].text = str(player11stat[31])
        p2text[497].text = str(player11stat[32])
        p2text[499].text = str(player11stat[33])

        p2text[503].text = ''

    # Team Stats

    if len(playerstats) == 11:

        teamtext[0].text = str(teamstats[1])
        teamtext[4].text = str(teamstats[3])
        teamtext[11].text = str(paintdist)
        teamtext[13].text = str(middist)
        teamtext[15].text = str(threedist)
        teamtext[19].text = str(teamstats[6])
        teamtext[22].text = str(teamstats[4])
        teamtext[24].text = str(teamstats[5])
        teamtext[28].text = str(paintpct)
        teamtext[31].text = str(paintmg)
        teamtext[33].text = str(paintag)
        teamtext[37].text = str(midpct)
        teamtext[40].text = str(midmg)
        teamtext[42].text = str(midag)
        teamtext[46].text = str(teamstats[9])
        teamtext[49].text = str(teamstats[7])
        teamtext[51].text = str(teamstats[8])
        teamtext[55].text = str(teamstats[12])
        teamtext[58].text = str(teamstats[10])
        teamtext[60].text = str(teamstats[11])
        teamtext[64].text = str(teamstats[15])
        teamtext[66].text = str(teamstats[13])
        teamtext[68].text = str(teamstats[14])
        teamtext[70].text = str(teamstats[16])
        teamtext[72].text = str(teamstats[19])
        teamtext[75].text = str(teamstats[17])
        teamtext[77].text = str(teamstats[18])

        teamtext[82].text = str(oppstats[1])
        teamtext[86].text = str(oppstats[4])
        teamtext[89].text = str(oppstats[2])
        teamtext[91].text = str(oppstats[3])
        teamtext[95].text = str(oppstats[7])
        teamtext[98].text = str(oppstats[5])
        teamtext[100].text = str(oppstats[6])
        teamtext[104].text = str(oppstats[10])
        teamtext[107].text = str(oppstats[8])
        teamtext[109].text = str(oppstats[9])
        teamtext[113].text = str(oppstats[13])
        teamtext[115].text = str(oppstats[11])
        teamtext[117].text = str(oppstats[12])
        teamtext[119].text = str(oppstats[14])
        teamtext[121].text = str(oppstats[17])
        teamtext[124].text = str(oppstats[15])
        teamtext[126].text = str(oppstats[16])


    if len(playerstats) < 11:

        teamtext[0].text = str(teamstats[1])
        teamtext[4].text = str(teamstats[3])
        teamtext[11].text = str(paintdist)
        teamtext[13].text = str(middist)
        teamtext[15].text = str(threedist)
        teamtext[18].text = str(teamstats[6])
        teamtext[21].text = str(teamstats[4])
        teamtext[23].text = str(teamstats[5])
        teamtext[27].text = str(paintpct)
        teamtext[30].text = str(paintmg)
        teamtext[32].text = str(paintag)
        teamtext[36].text = str(midpct)
        teamtext[39].text = str(midmg)
        teamtext[41].text = str(midag)
        teamtext[44].text = str(teamstats[9])
        teamtext[47].text = str(teamstats[7])
        teamtext[49].text = str(teamstats[8])
        teamtext[53].text = str(teamstats[12])
        teamtext[56].text = str(teamstats[10])
        teamtext[58].text = str(teamstats[11])
        teamtext[62].text = str(teamstats[15])
        teamtext[64].text = str(teamstats[13])
        teamtext[66].text = str(teamstats[14])
        teamtext[68].text = str(teamstats[16])
        teamtext[70].text = str(teamstats[19])
        teamtext[73].text = str(teamstats[17])
        teamtext[75].text = str(teamstats[18])

        teamtext[80].text = str(oppstats[1])
        teamtext[83].text = str(oppstats[4])
        teamtext[86].text = str(oppstats[2])
        teamtext[88].text = str(oppstats[3])
        teamtext[91].text = str(oppstats[7])
        teamtext[94].text = str(oppstats[5])
        teamtext[96].text = str(oppstats[6])
        teamtext[100].text = str(oppstats[10])
        teamtext[103].text = str(oppstats[8])
        teamtext[105].text = str(oppstats[9])
        teamtext[109].text = str(oppstats[13])
        teamtext[111].text = str(oppstats[11])
        teamtext[113].text = str(oppstats[12])
        teamtext[115].text = str(oppstats[14])
        teamtext[117].text = str(oppstats[17])
        teamtext[120].text = str(oppstats[15])
        teamtext[122].text = str(oppstats[16])

    teamname = str(scrape.getteamnamefromabbr(scoutteam))

    #document.save(teamname + 'requestedscout.docx')

    return document

